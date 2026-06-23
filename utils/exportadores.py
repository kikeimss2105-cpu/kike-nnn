from io import BytesIO
from datetime import datetime

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def aplicar_borde_tabla(tabla):
    tbl = tabla._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        borders.append(border)

    tbl_pr.append(borders)


def sombrear_celda(celda, color="D9EAF7"):
    tc_pr = celda._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    tc_pr.append(shading)


def texto_negritas(celda):
    for parrafo in celda.paragraphs:
        for run in parrafo.runs:
            run.bold = True


def configurar_celda(celda):
    celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for parrafo in celda.paragraphs:
        for run in parrafo.runs:
            run.font.size = Pt(9)


def dato(fila, columna, default=""):
    try:
        valor = fila.get(columna, default)
        if pd.isna(valor):
            return default
        return valor
    except Exception:
        return default


def agregar_tabla_diccionario(document, titulo, datos):
    document.add_heading(titulo, level=2)

    tabla = document.add_table(rows=1, cols=2)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.style = "Table Grid"

    encabezados = tabla.rows[0].cells
    encabezados[0].text = "Campo"
    encabezados[1].text = "Información"

    for celda in encabezados:
        sombrear_celda(celda)
        texto_negritas(celda)
        configurar_celda(celda)

    for clave, valor in datos.items():
        fila = tabla.add_row().cells
        fila[0].text = str(clave)
        fila[1].text = str(valor)
        configurar_celda(fila[0])
        configurar_celda(fila[1])

    aplicar_borde_tabla(tabla)
    document.add_paragraph("")


def agregar_tabla_resumen_diagnosticos(document, df_resultados):
    document.add_heading("3. Diagnósticos sugeridos", level=2)
    document.add_paragraph(
        "Tabla compacta para identificar rápidamente diagnóstico, puntaje, confianza, jerarquía y prioridad."
    )

    columnas = ["Código", "NANDA", "Puntaje", "Confianza", "Jerarquía", "Prioridad"]
    tabla = document.add_table(rows=1, cols=len(columnas))
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.style = "Table Grid"

    for i, columna in enumerate(columnas):
        celda = tabla.rows[0].cells[i]
        celda.text = columna
        sombrear_celda(celda)
        texto_negritas(celda)
        configurar_celda(celda)

    for _, fila_df in df_resultados.iterrows():
        fila = tabla.add_row().cells
        for i, columna in enumerate(columnas):
            fila[i].text = str(dato(fila_df, columna))
            configurar_celda(fila[i])

    aplicar_borde_tabla(tabla)
    document.add_paragraph("")


def agregar_tabla_vinculacion_nnn(document, df_resultados):
    document.add_heading("4. Vinculación NANDA-NOC-NIC", level=2)
    document.add_paragraph(
        "Tabla separada para evitar columnas saturadas y mejorar la lectura clínica."
    )

    columnas = ["NANDA", "Coincidencias clínicas", "NOC sugerido", "NIC sugerido"]
    tabla = document.add_table(rows=1, cols=len(columnas))
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.style = "Table Grid"

    for i, columna in enumerate(columnas):
        celda = tabla.rows[0].cells[i]
        celda.text = columna
        sombrear_celda(celda)
        texto_negritas(celda)
        configurar_celda(celda)

    for _, fila_df in df_resultados.iterrows():
        fila = tabla.add_row().cells
        fila[0].text = str(dato(fila_df, "NANDA"))
        fila[1].text = str(dato(fila_df, "Coincidencias"))
        fila[2].text = str(dato(fila_df, "NOC sugerido"))
        fila[3].text = str(dato(fila_df, "NIC sugerido"))

        for celda in fila:
            configurar_celda(celda)

    aplicar_borde_tabla(tabla)
    document.add_paragraph("")


def agregar_plan_narrativo(document, df_resultados):
    document.add_heading("5. Plan narrativo NANDA-NOC-NIC", level=2)

    for _, fila in df_resultados.iterrows():
        document.add_heading(str(dato(fila, "NANDA")), level=3)

        datos = {
            "Código": dato(fila, "Código"),
            "Dominio": dato(fila, "Dominio"),
            "Clase": dato(fila, "Clase"),
            "Definición": dato(fila, "Definición"),
            "Coincidencias clínicas": dato(fila, "Coincidencias"),
            "Puntaje": dato(fila, "Puntaje"),
            "Confianza educativa": dato(fila, "Confianza"),
            "Jerarquía": dato(fila, "Jerarquía"),
            "NOC sugerido": dato(fila, "NOC sugerido"),
            "NIC sugerido": dato(fila, "NIC sugerido"),
            "Prioridad": dato(fila, "Prioridad"),
            "Meta esperada": dato(fila, "Meta esperada"),
            "Indicadores NOC": dato(fila, "Indicadores NOC"),
            "Actividades NIC": dato(fila, "Actividades NIC"),
            "Fundamentos": dato(fila, "Fundamentos"),
            "Nota": dato(fila, "Nota", "Requiere validación clínica"),
        }

        agregar_tabla_diccionario(document, "Detalle del diagnóstico", datos)


def generar_excel(df_resultados, datos_paciente):
    output = BytesIO()

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        pd.DataFrame([datos_paciente]).to_excel(
            writer,
            index=False,
            sheet_name="Datos del paciente"
        )

        columnas_resumen = [
            col for col in ["Código", "NANDA", "Puntaje", "Confianza", "Jerarquía", "Prioridad"]
            if col in df_resultados.columns
        ]

        if columnas_resumen:
            df_resultados[columnas_resumen].to_excel(
                writer,
                index=False,
                sheet_name="Resumen diagnóstico"
            )

        df_resultados.to_excel(
            writer,
            index=False,
            sheet_name="Plan NNN completo"
        )

        glosario = pd.DataFrame([
            {"Elemento": "NANDA", "Descripción": "Taxonomía diagnóstica de enfermería."},
            {"Elemento": "NOC", "Descripción": "Clasificación de resultados esperados de enfermería."},
            {"Elemento": "NIC", "Descripción": "Clasificación de intervenciones de enfermería."},
            {"Elemento": "Braden", "Descripción": "Escala para riesgo de lesiones por presión."},
            {"Elemento": "EVA", "Descripción": "Escala visual analógica para dolor."},
            {"Elemento": "Glasgow", "Descripción": "Escala de respuesta neurológica."},
            {"Elemento": "Advertencia", "Descripción": "Uso educativo. No sustituye juicio profesional."},
        ])

        glosario.to_excel(writer, index=False, sheet_name="Glosario")

        workbook = writer.book
        for sheet_name in workbook.sheetnames:
            hoja = workbook[sheet_name]
            for columna in hoja.columns:
                longitud_maxima = 0
                letra_columna = columna[0].column_letter
                for celda in columna:
                    try:
                        longitud_maxima = max(longitud_maxima, len(str(celda.value)))
                    except Exception:
                        pass
                hoja.column_dimensions[letra_columna].width = min(longitud_maxima + 3, 55)

    output.seek(0)
    return output


def generar_word(df_resultados, datos_paciente):
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    titulo = document.add_heading("Plan de Cuidados NANDA-NIC-NOC", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitulo = document.add_paragraph(
        "Sistema KIKE-NNN | Documento educativo generado automáticamente"
    )
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fecha = datetime.now().strftime("%d/%m/%Y %H:%M")
    parrafo_fecha = document.add_paragraph(f"Fecha de generación: {fecha}")
    parrafo_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")
    document.add_paragraph(
        "Nota de seguridad: Este documento tiene finalidad educativa. "
        "Las sugerencias NANDA-NOC-NIC requieren validación clínica, valoración integral "
        "del paciente y apego a protocolos institucionales."
    )

    agregar_tabla_diccionario(document, "1. Datos del paciente", datos_paciente)

    if df_resultados.empty:
        document.add_heading("2. Resumen del análisis", level=2)
        document.add_paragraph("No se encontraron diagnósticos sugeridos.")
        output = BytesIO()
        document.save(output)
        output.seek(0)
        return output

    resumen = {
        "Número de diagnósticos sugeridos": len(df_resultados),
        "Diagnóstico con mayor puntaje": df_resultados.iloc[0].get("NANDA", ""),
        "Puntaje más alto": df_resultados.iloc[0].get("Puntaje", ""),
        "Confianza educativa máxima": df_resultados.iloc[0].get("Confianza", ""),
        "Jerarquía del diagnóstico principal": df_resultados.iloc[0].get("Jerarquía", ""),
    }

    agregar_tabla_diccionario(document, "2. Resumen del análisis", resumen)
    agregar_tabla_resumen_diagnosticos(document, df_resultados)
    agregar_tabla_vinculacion_nnn(document, df_resultados)
    agregar_plan_narrativo(document, df_resultados)

    document.add_heading("6. Recomendación académica", level=2)
    document.add_paragraph(
        "El estudiante debe contrastar los diagnósticos sugeridos con la valoración completa, "
        "la taxonomía NANDA-I vigente, las clasificaciones NOC/NIC, el expediente clínico, "
        "los signos vitales, estudios complementarios y el contexto individual del paciente."
    )

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output
